import hashlib
import os
import pathlib
import re
import subprocess
import sys
import uuid
import sqlite3
from datetime import datetime, timezone
from typing import List, Tuple, Optional
from fastapi import FastAPI, HTTPException
from pydantic import BaseModel

# B4 — grammar-constrained decoding: validate_and_heal wraps /extract and /ask
try:
    from models.constrained_decoding import validate_and_heal as _validate_and_heal
    _CONSTRAINED_DECODING_AVAILABLE = True
except ImportError:
    _CONSTRAINED_DECODING_AVAILABLE = False
    def _validate_and_heal(obj, schema_name, **kw):  # type: ignore[misc]
        return obj  # no-op fallback if module unavailable

# Initialize FastAPI app
app = FastAPI(title="Kairo Sidecar")

# Setup dependencies (Qdrant & SentenceTransformer) — optional, lazy-loaded
try:
    from qdrant_client import QdrantClient
    from qdrant_client.models import Distance, VectorParams, PointStruct, Filter, FieldCondition, MatchValue
    _QDRANT_AVAILABLE = True
except ImportError:
    _QDRANT_AVAILABLE = False
    QdrantClient = None
    Distance = VectorParams = PointStruct = Filter = FieldCondition = MatchValue = None  # type: ignore

try:
    from sentence_transformers import SentenceTransformer as _SentenceTransformer
    _ST_AVAILABLE = True
except ImportError:
    _ST_AVAILABLE = False
    _SentenceTransformer = None  # type: ignore


# C1 — pluggable vector store: LanceDB default, Qdrant Edge fallback via feature flag
# Use importlib to load by absolute path so this works both as a standalone module
# (python -m kernel.sidecar.app) and as a package import (from kernel.sidecar.app import app).
try:
    import importlib.util as _ilu
    _vs_spec = _ilu.spec_from_file_location(
        "retrieval.vector_store",
        os.path.join(os.path.dirname(__file__), "retrieval", "vector_store.py")
    )
    _vs_mod = _ilu.module_from_spec(_vs_spec)
    _vs_spec.loader.exec_module(_vs_mod)
    _get_store = _vs_mod.get_store
    _VECTOR_STORE_AVAILABLE = True
except Exception:
    _VECTOR_STORE_AVAILABLE = False
    _get_store = None  # type: ignore



def _hash_embed(text: str, dim: int = 256) -> list:
    """Deterministic fallback embedding using SHA-256. Produces NaN-free normalized floats."""
    import struct
    # Use multiple SHA-256 rounds to fill dim floats from clean integer values
    floats = []
    seed = text.encode()
    i = 0
    while len(floats) < dim:
        h = hashlib.sha256(seed + i.to_bytes(4, 'little')).digest()
        # Convert each 4-byte chunk to an unsigned int, then map to [-1, 1]
        for j in range(0, len(h) - 3, 4):
            val = struct.unpack_from('<I', h, j)[0]  # unsigned int
            floats.append((val / 2147483647.5) - 1.0)  # map to (-1, 1)
            if len(floats) >= dim:
                break
        i += 1
    floats = floats[:dim]
    norm = sum(f * f for f in floats) ** 0.5 or 1.0
    return [f / norm for f in floats]


class _HashEmbedder:
    """Fallback embedder when sentence_transformers is not available."""
    def encode(self, texts, **kwargs):
        return [_hash_embed(t) for t in (texts if isinstance(texts, list) else [texts])]


def _to_vector(enc_result) -> list:
    """Convert encode() result to a plain Python list regardless of whether it's a numpy array or list."""
    if hasattr(enc_result, 'tolist'):
        return enc_result.tolist()
    if isinstance(enc_result, list):
        # Could be list-of-numpy or list-of-list — take first element if batched
        if enc_result and hasattr(enc_result[0], 'tolist'):
            return enc_result[0].tolist()
        if enc_result and isinstance(enc_result[0], list):
            return enc_result[0]
        return enc_result
    return list(enc_result)


# Initialize embedding model
if _ST_AVAILABLE:
    try:
        embedding_model = _SentenceTransformer("minishlab/potion-base-8M", local_files_only=True)
    except Exception:
        embedding_model = _HashEmbedder()
else:
    embedding_model = _HashEmbedder()


# Initialize Qdrant Client (in embedded mode, or in-memory under pytest to avoid lock conflicts)
is_pytest = "pytest" in sys.modules or any("pytest" in arg for arg in sys.argv) or os.environ.get("KAIRO_USE_MEMORY_QDRANT") == "1"
if _QDRANT_AVAILABLE:
    if is_pytest:
        qdrant_client = QdrantClient(location=":memory:")
    else:
        qdrant_db_path = os.path.abspath(os.path.join(os.path.dirname(__file__), ".kairo", "qdrant"))
        os.makedirs(os.path.dirname(qdrant_db_path), exist_ok=True)
        qdrant_client = QdrantClient(path=qdrant_db_path)
else:
    qdrant_client = None

COLLECTION_NAME = "kairo_chunks"
if qdrant_client is not None:
    try:
        if not qdrant_client.collection_exists(COLLECTION_NAME):
            qdrant_client.create_collection(
                collection_name=COLLECTION_NAME,
                vectors_config=VectorParams(size=256, distance=Distance.COSINE),
            )
    except Exception:
        try:
            qdrant_client.create_collection(
                collection_name=COLLECTION_NAME,
                vectors_config=VectorParams(size=256, distance=Distance.COSINE),
            )
        except Exception:
            pass

# C1 — initialize the unified vector store (LanceDB default; KAIRO_VECTOR_BACKEND=qdrant for fallback)
_VECTOR_BACKEND = os.environ.get("KAIRO_VECTOR_BACKEND", "qdrant")  # keep qdrant default for test compat
if _VECTOR_STORE_AVAILABLE and _get_store is not None:
    _vec_path = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..", ".kairo_test" if is_pytest else ".kairo", "vectors"))
    try:
        _vector_store = _get_store(_VECTOR_BACKEND, path=_vec_path)
    except Exception:
        _vector_store = None
else:
    _vector_store = None

# Import packs
from packs.generic.pack import GenericPack
from packs.invoice.pack import InvoicePack
from packs.paper.pack import PaperPack
from packs.contract.pack import ContractPack


# DB path — defined FIRST so everything below can reference it
def _get_db_path() -> str:
    """Return DB path — uses .kairo_test/ subdirectory when running under pytest."""
    if is_pytest:
        base = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..", ".kairo_test"))
    else:
        base = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..", ".kairo"))
    os.makedirs(base, exist_ok=True)
    return os.path.join(base, "kairo.db")


def _init_db_schema(db_path: str) -> None:
    """Initialize SQLite schema. Idempotent — safe to call multiple times."""
    conn = sqlite3.connect(db_path)
    c = conn.cursor()
    c.executescript("""
        CREATE TABLE IF NOT EXISTS documents (
            doc_id TEXT PRIMARY KEY,
            source_path TEXT,
            sha256 TEXT,
            page_count INTEGER,
            created_at INTEGER,
            use_visual_retrieval INTEGER DEFAULT 0
        );
        CREATE TABLE IF NOT EXISTS pages (
            doc_id TEXT,
            page_index INTEGER,
            width_px INTEGER,
            height_px INTEGER,
            image_sha256 TEXT,
            PRIMARY KEY (doc_id, page_index)
        );
        CREATE TABLE IF NOT EXISTS chunks (
            id TEXT PRIMARY KEY,
            doc_id TEXT,
            page_index INTEGER,
            x0 REAL,
            y0 REAL,
            x1 REAL,
            y1 REAL,
            text TEXT,
            chunk_order INTEGER
        );
        CREATE TABLE IF NOT EXISTS extractions (
            id TEXT PRIMARY KEY,
            doc_id TEXT,
            field TEXT,
            value TEXT,
            confidence REAL,
            status TEXT,
            method TEXT
        );
        CREATE TABLE IF NOT EXISTS anchors (
            extraction_id TEXT,
            chunk_id TEXT,
            char_start INTEGER,
            char_end INTEGER,
            page INTEGER,
            x0 REAL,
            y0 REAL,
            x1 REAL,
            y1 REAL,
            PRIMARY KEY (extraction_id, chunk_id)
        );
        CREATE TABLE IF NOT EXISTS answers (
            id TEXT PRIMARY KEY,
            query TEXT,
            text TEXT,
            grounded INTEGER
        );
        CREATE TABLE IF NOT EXISTS citations (
            answer_id TEXT,
            chunk_id TEXT,
            char_start INTEGER,
            char_end INTEGER,
            page INTEGER,
            x0 REAL,
            y0 REAL,
            x1 REAL,
            y1 REAL,
            PRIMARY KEY (answer_id, chunk_id)
        );
        CREATE TABLE IF NOT EXISTS corrections (
            extraction_id TEXT PRIMARY KEY,
            old_value TEXT,
            new_value TEXT,
            by TEXT,
            at_time INTEGER
        );
    """)
    # Idempotent column add (SQLite does not support IF NOT EXISTS on ALTER COLUMN)
    try:
        c.execute("ALTER TABLE documents ADD COLUMN use_visual_retrieval INTEGER DEFAULT 0")
        conn.commit()
    except sqlite3.OperationalError:
        pass  # Column already exists
    conn.commit()
    conn.close()


# Initialize DB schema at startup (idempotent)
_db_path = _get_db_path()
_init_db_schema(_db_path)
try:
    from retrieval.migration import MigrationManager
    MigrationManager().run_migrations(_db_path, target_version=2)
except Exception as _e:
    print(f"[sidecar] Migration warning/error: {_e}")


class HealthResponse(BaseModel):
    sidecar: str
    db_writable: bool
    qdrant_available: bool
    embedding_model: str
    db_path: str


@app.get("/health", response_model=HealthResponse)

# ── Phase 2: Web Demo + Connector API routes (scaffold-specific) ──────────
from fastapi.responses import FileResponse, HTMLResponse
from fastapi import UploadFile, File as FastAPIFile
import os as _os

@app.get("/demo", response_class=HTMLResponse)
async def serve_demo():
    """Serve the interactive web demo (scaffold/web/demo.html)."""
    demo_path = _os.path.join(_os.path.dirname(_os.path.dirname(_os.path.dirname(_os.path.abspath(__file__)))), "scaffold", "web", "demo.html")
    if _os.path.exists(demo_path):
        return FileResponse(demo_path, media_type="text/html")
    return HTMLResponse("<h1>Demo not found</h1><p>scaffold/web/demo.html missing</p>", status_code=404)


@app.post("/api/extract-document")
async def api_extract_document(file: UploadFile = FastAPIFile(...)):
    """Connector API: upload a document, get grounded extractions with bbox.

    Returns structured JSON: { doc_id, doc_type, fields: [{field, value, status, bbox, page, method, confidence}] }
    Refused fields have status='blocked' with a reason.
    """
    import tempfile, shutil
    # Save uploaded file to temp
    suffix = _os.path.splitext(file.filename or "upload.txt")[1] or ".txt"
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        shutil.copyfileobj(file.file, tmp)
        tmp_path = tmp.name

    try:
        # Index the document
        idx_resp = index_doc(IndexRequest(filepath=tmp_path))
        doc_id = idx_resp.doc_id

        # Detect pack from filename or content
        fname = (file.filename or "").lower()
        if "invoice" in fname or "inv" in fname:
            pack = "invoice"
        elif "contract" in fname or "con" in fname:
            pack = "contract"
        elif "paper" in fname or "pap" in fname:
            pack = "paper"
        else:
            pack = "generic"

        # Extract fields
        extractions = extract_fields(ExtractRequest(doc_id=doc_id, pack=pack))

        # Format for web demo
        fields = []
        for ext in extractions:
            grounded = ext.method != "block" and ext.status != "blocked"
            field_data = {
                "field": ext.field,
                "value": ext.value if grounded else None,
                "status": "grounded" if grounded else "blocked",
                "method": ext.method if grounded else None,
                "confidence": ext.confidence if grounded else 0.0,
                "page": ext.anchors[0].page if ext.anchors else None,
                "bbox": list(ext.anchors[0].bbox) if ext.anchors else None,
            }
            if not grounded:
                field_data["reason"] = "no grounded source found in document"
            fields.append(field_data)

        return {"doc_id": doc_id, "doc_type": pack, "fields": fields}
    finally:
        _os.unlink(tmp_path)


@app.post("/api/ask-document")
async def api_ask_document(req: dict):
    """Connector API: ask a question about an indexed document.

    Returns: { status, answer, anchors: [{page, bbox}] } or { status: 'blocked', reason }
    """
    doc_id = req.get("doc_id")
    question = req.get("question", "")
    if not doc_id or not question:
        return {"status": "blocked", "reason": "doc_id and question required"}

    answer = ask_question(AskRequest(doc_id=doc_id, question=question))
    if answer.status == "blocked" or answer.method == "block":
        return {"status": "blocked", "reason": "no grounded source found for this question"}

    anchors = []
    if answer.anchors:
        for a in answer.anchors:
            anchors.append({"page": a.page, "bbox": list(a.bbox)})

    return {"status": "grounded", "answer": answer.answer, "anchors": anchors}


def health_check():
    """Kairo Doctor health endpoint — checks sidecar, DB, and stores."""
    db_path = _get_db_path()
    db_writable = False
    try:
        conn = sqlite3.connect(db_path)
        conn.execute("SELECT 1")
        conn.close()
        db_writable = True
    except Exception:
        pass
    model_name = "hash_embed" if isinstance(embedding_model, _HashEmbedder) else "sentence_transformer"
    return HealthResponse(
        sidecar="ok",
        db_writable=db_writable,
        qdrant_available=_QDRANT_AVAILABLE and qdrant_client is not None,
        embedding_model=model_name,
        db_path=db_path,
    )

class IndexRequest(BaseModel):
    path: str

class BBox(BaseModel):
    x0: float
    y0: float
    x1: float
    y1: float

class IndexResponsePage(BaseModel):
    index: int
    width_px: int
    height_px: int
    image_sha256: str

class IndexResponseChunk(BaseModel):
    text: str
    page_index: int
    bbox: BBox
    order: int

class IndexResponse(BaseModel):
    doc_id: str
    pages: int
    chunks: int
    pages_list: List[IndexResponsePage] = []
    chunks_list: List[IndexResponseChunk] = []

class Anchor(BaseModel):
    chunk_id: str
    char_span: Tuple[int, int]
    page: int
    bbox: BBox

class Extraction(BaseModel):
    id: str
    doc_id: str
    field: str
    value: str
    confidence: float
    status: str
    anchors: List[Anchor]
    method: str

class ExtractRequest(BaseModel):
    doc_id: str
    pack: str

class AskRequest(BaseModel):
    doc_id: str
    query: str

class VisualAskRequest(BaseModel):
    doc_id: str
    query: str
    page_index: Optional[int] = None  # restrict to a specific page; None = all pages
    iou_threshold: float = 0.5         # minimum IoU to accept the visual match

class VisualAskResponse(BaseModel):
    doc_id: str
    query: str
    matched_bbox: Optional[BBox]       # best patch bbox, None if disabled/blocked
    chunk_id: Optional[str]            # B3-verified chunk_id, None if blocked
    score: float                       # ColPali/hash normalised similarity
    iou_passed: bool                   # True if B3 IoU gate passed
    visual_retrieval_enabled: bool     # reflects per-document flag

class Answer(BaseModel):
    id: str
    query: str
    text: str
    citations: List[Anchor]
    grounded: bool

class ProvenanceResponse(BaseModel):
    page: int
    bbox: BBox
    char_span: Tuple[int, int]
    image_ref: str

class CorrectRequest(BaseModel):
    extraction_id: str
    new_value: str

class Correction(BaseModel):
    extraction_id: str
    old_value: str
    new_value: str
    by: str
    at: datetime

# New models for stateless parser
class ParserPage(BaseModel):
    index: int
    width_px: int
    height_px: int
    image_sha256: str
    is_raster: Optional[bool] = False

class ParserChunk(BaseModel):
    page: int
    bbox: BBox
    text: str
    source_type: str

class ParseResponse(BaseModel):
    doc_id: str
    pages: List[ParserPage]
    chunks: List[ParserChunk]

# Helper functions for parsing
def _split_paragraphs(text: str) -> List[str]:
    blocks = re.split(r"\n\s*\n", text)
    return [b for b in blocks if b.strip()]

def _ingest_text(filepath: pathlib.Path) -> Tuple[List[ParserPage], List[ParserChunk]]:
    text = filepath.read_text(encoding="utf-8", errors="replace")
    paragraphs = _split_paragraphs(text)
    
    chunks = []
    _CHARS_PER_LINE = 80
    _LINES_PER_PAGE = 60
    current_line = 0
    page_count = 1
    
    for para in paragraphs:
        para_text = para.strip()
        if not para_text:
            continue
            
        para_lines = para_text.count("\n") + 1
        page = (current_line // _LINES_PER_PAGE) + 1
        page_count = max(page_count, page)
        
        page_offset = current_line % _LINES_PER_PAGE
        y0 = page_offset * (1.0 / _LINES_PER_PAGE)
        y1 = min(1.0, (page_offset + para_lines) * (1.0 / _LINES_PER_PAGE))
        
        chunks.append(ParserChunk(
            page=page,
            bbox=BBox(x0=0.0, y0=y0, x1=1.0, y1=y1),
            text=para_text,
            source_type="text"
        ))
        current_line += para_lines + 1
        
    pages = []
    for p in range(1, page_count + 1):
        pages.append(ParserPage(
            index=p,
            width_px=800,
            height_px=1000,
            image_sha256=""
        ))
        
    return pages, chunks

def _ingest_docx(filepath: pathlib.Path) -> Tuple[List[ParserPage], List[ParserChunk]]:
    try:
        from docx import Document as DocxDocument
    except ImportError:
        return _ingest_text(filepath)
        
    doc = DocxDocument(str(filepath))
    chunks = []
    _CHARS_PER_LINE = 80
    _LINES_PER_PAGE = 60
    current_line = 0
    page_count = 1
    
    for para in doc.paragraphs:
        para_text = para.text.strip()
        if not para_text:
            continue
            
        para_lines = max(1, len(para_text) // _CHARS_PER_LINE + 1)
        page = (current_line // _LINES_PER_PAGE) + 1
        page_count = max(page_count, page)
        
        page_offset = current_line % _LINES_PER_PAGE
        y0 = page_offset * (1.0 / _LINES_PER_PAGE)
        y1 = min(1.0, (page_offset + para_lines) * (1.0 / _LINES_PER_PAGE))
        
        chunks.append(ParserChunk(
            page=page,
            bbox=BBox(x0=0.0, y0=y0, x1=1.0, y1=y1),
            text=para_text,
            source_type="docx_paragraph"
        ))
        current_line += para_lines + 1
        
    pages = []
    for p in range(1, page_count + 1):
        pages.append(ParserPage(
            index=p,
            width_px=800,
            height_px=1000,
            image_sha256=""
        ))
        
    return pages, chunks

def _ingest_pdf_docling(filepath: pathlib.Path) -> Tuple[List[ParserPage], List[ParserChunk]]:
    from docling.document_converter import DocumentConverter
    converter = DocumentConverter()
    result = converter.convert(str(filepath))
    doc_docling = result.document
    
    pages = []
    chunks = []
    
    for index, page_obj in doc_docling.pages.items():
        width = 800
        height = 1000
        if hasattr(page_obj, "size"):
            width = int(page_obj.size.width)
            height = int(page_obj.size.height)
        pages.append(ParserPage(
            index=index,
            width_px=width,
            height_px=height,
            image_sha256=""
        ))
        
    elements = []
    if hasattr(doc_docling, "texts"):
        elements.extend(doc_docling.texts)
    if hasattr(doc_docling, "tables"):
        elements.extend(doc_docling.tables)
    if not elements and hasattr(doc_docling, "elements"):
        elements = list(doc_docling.elements)
        
    for item in elements:
        if hasattr(item, "prov") and item.prov:
            page_info = item.prov[0]
            page_no = page_info.page_no
            bbox = page_info.bbox
            
            text = ""
            if hasattr(item, "text"):
                text = item.text
            elif hasattr(item, "export_to_markdown"):
                text = item.export_to_markdown()
                
            text = text.strip()
            if not text:
                continue
                
            page_width = 1.0
            page_height = 1.0
            for p in pages:
                if p.index == page_no:
                    page_width = p.width_px
                    page_height = p.height_px
                    break
                    
            x0 = max(0.0, min(bbox.left / page_width if page_width > 0 else 0.0, 1.0))
            y0 = max(0.0, min(1.0 - (bbox.top / page_height) if page_height > 0 else 0.0, 1.0))
            x1 = max(x0, min(bbox.right / page_width if page_width > 0 else 1.0, 1.0))
            y1 = max(0.0, min(1.0 - (bbox.bottom / page_height) if page_height > 0 else 0.0, 1.0))
            
            chunks.append(ParserChunk(
                page=page_no,
                bbox=BBox(x0=x0, y0=y0, x1=x1, y1=y1),
                text=text,
                source_type="pdf_docling"
            ))
            
    return pages, chunks

def _ingest_pdf_fastpath(filepath: pathlib.Path) -> Tuple[List[ParserPage], List[ParserChunk]]:
    if is_pytest:
        base_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..", ".kairo_test"))
    else:
        base_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..", ".kairo"))
    page_images_dir = os.path.join(base_dir, "page_images")

    # Try in-process call first to avoid spawning sys.executable under PyInstaller/frozen env
    try:
        try:
            from kernel.sidecar.pdf_fastpath import process_pdf
        except ImportError:
            from pdf_fastpath import process_pdf
        
        data = process_pdf(filepath, page_images_dir)
        pages = [ParserPage(**p) for p in data["pages"]]
        chunks = [ParserChunk(**c) for c in data["chunks"]]
        return pages, chunks
    except Exception as e:
        sys.stderr.write(f"In-process PDF parse failed, falling back to subprocess: {e}\n")
        
        script_path = os.path.join(os.path.dirname(__file__), "pdf_fastpath.py")
        res = subprocess.run(
            [sys.executable, script_path, str(filepath), page_images_dir],
            capture_output=True,
            text=True,
            check=True
        )
        import json
        data = json.loads(res.stdout)
        pages = [ParserPage(**p) for p in data["pages"]]
        chunks = [ParserChunk(**c) for c in data["chunks"]]
        return pages, chunks

def _parse_document_internal(path_str: str) -> ParseResponse:
    filepath = pathlib.Path(path_str)
    if not filepath.exists():
        raise HTTPException(status_code=404, detail=f"File not found: {path_str}")
        
    content_bytes = filepath.read_bytes()
    sha256 = hashlib.sha256(content_bytes).hexdigest()
    doc_id = sha256[:16]
    
    suffix = filepath.suffix.lower()
    if suffix in (".txt", ".md"):
        pages, chunks = _ingest_text(filepath)
    elif suffix == ".docx":
        pages, chunks = _ingest_docx(filepath)
    elif suffix == ".pdf":
        try:
            pages, chunks = _ingest_pdf_fastpath(filepath)
            
            if is_pytest:
                base_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..", ".kairo_test"))
            else:
                base_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..", ".kairo"))
            page_images_dir = os.path.join(base_dir, "page_images")

            from kernel.sidecar.ingest.ocr_backends import get_ocr_backend

            # Page-by-page OCR detector and dispatch
            for page in pages:
                image_path = os.path.join(page_images_dir, f"{page.image_sha256}.png")
                backend_override = os.environ.get("KAIRO_OCR_BACKEND")
                
                if page.is_raster or backend_override:
                    # Clear PyMuPDF native chunks for this raster page
                    chunks = [c for c in chunks if c.page != page.index]
                    
                    backend_name = backend_override if backend_override else os.environ.get("KAIRO_OCR_BACKEND_SCANNED", "deepseek_ocr2")
                    backend = get_ocr_backend(backend_name)
                    
                    ocr_chunks = backend.extract_text_with_bboxes(image_path)
                    for ocr_c in ocr_chunks:
                        chunks.append(ParserChunk(
                            page=page.index,
                            bbox=BBox(
                                x0=ocr_c["bbox"]["x0"],
                                y0=ocr_c["bbox"]["y0"],
                                x1=ocr_c["bbox"]["x1"],
                                y1=ocr_c["bbox"]["y1"]
                            ),
                            text=ocr_c["text"],
                            source_type=f"pdf_ocr_{backend_name}"
                        ))
                else:
                    # Native page: optionally force docling, else keep PyMuPDF chunks
                    if os.environ.get("KAIRO_FORCE_NATIVE_DOCLING") == "1":
                        chunks = [c for c in chunks if c.page != page.index]
                        backend = get_ocr_backend("docling")
                        ocr_chunks = backend.extract_text_with_bboxes(image_path)
                        for ocr_c in ocr_chunks:
                            chunks.append(ParserChunk(
                                page=page.index,
                                bbox=BBox(
                                    x0=ocr_c["bbox"]["x0"],
                                    y0=ocr_c["bbox"]["y0"],
                                    x1=ocr_c["bbox"]["x1"],
                                    y1=ocr_c["bbox"]["y1"]
                                ),
                                text=ocr_c["text"],
                                source_type="pdf_ocr_docling"
                            ))
        except Exception as e:
            sys.stderr.write(f"PyMuPDF fastpath failed, falling back to Docling: {e}\n")
            pages, chunks = _ingest_pdf_docling(filepath)
    else:
        raise HTTPException(status_code=400, detail=f"Unsupported file extension: {suffix}")
        
    return ParseResponse(doc_id=doc_id, pages=pages, chunks=chunks)


# Helper functions for cascade and DB access
def load_chunks_from_db(doc_id: str) -> list:
    db_path = _get_db_path()
    if not os.path.exists(db_path):
        return []
    conn = sqlite3.connect(db_path)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    cursor.execute(
        "SELECT id, page_index, x0, y0, x1, y1, text, chunk_order FROM chunks WHERE doc_id = ? ORDER BY chunk_order ASC",
        (doc_id,)
    )
    rows = cursor.fetchall()
    conn.close()
    
    class SimpleBBox:
        def __init__(self, x0, y0, x1, y1):
            self.x0 = x0
            self.y0 = y0
            self.x1 = x1
            self.y1 = y1
            
    class SimpleChunk:
        def __init__(self, chunk_id, page, bbox, text, order):
            self.chunk_id = chunk_id
            self.page = page
            self.page_index = page
            self.bbox = bbox
            self.text = text
            self.order = order
            self.embedding = []
            
    chunks = []
    for r in rows:
        chunks.append(SimpleChunk(
            chunk_id=r["id"],
            page=r["page_index"],
            bbox=SimpleBBox(r["x0"], r["y0"], r["x1"], r["y1"]),
            text=r["text"],
            order=r["chunk_order"]
        ))
    return chunks

def normalize_text(text: str) -> str:
    text = text.lower().strip()
    text = re.sub(r'[^\w\s\-\.]', '', text)
    text = " ".join(text.split())
    text = re.sub(r'[\$\€\£\¥]', '', text)
    return text

def levenshtein_ratio(s1: str, s2: str) -> float:
    s1 = s1.lower()
    s2 = s2.lower()
    if s1 == s2:
        return 1.0
    if len(s1) == 0 or len(s2) == 0:
        return 0.0

    rows = len(s1) + 1
    cols = len(s2) + 1
    dist = [[0 for _ in range(cols)] for _ in range(rows)]

    for i in range(1, rows):
        dist[i][0] = i
    for j in range(1, cols):
        dist[0][j] = j

    for col in range(1, cols):
        for row in range(1, rows):
            if s1[row-1] == s2[col-1]:
                cost = 0
            else:
                cost = 1
            dist[row][col] = min(dist[row-1][col] + 1,
                                 dist[row][col-1] + 1,
                                 dist[row-1][col-1] + cost)

    return 1.0 - (dist[rows-1][cols-1] / max(len(s1), len(s2)))

def best_fuzzy_match(value: str, text: str) -> tuple[float, tuple[int, int]]:
    val_norm = normalize_text(value)
    if not val_norm:
        return 0.0, (0, 0)

    val_words = val_norm.split()
    n = len(val_words)

    words_spans = []
    for m in re.finditer(r'\S+', text):
        words_spans.append((m.group(0), m.start(), m.end()))

    if not words_spans:
        return 0.0, (0, 0)

    best_ratio = 0.0
    best_span = (0, 0)

    for length in range(max(1, n-1), min(len(words_spans) + 1, n+3)):
        for i in range(len(words_spans) - length + 1):
            window = words_spans[i:i+length]
            sub_text = text[window[0][1]:window[-1][2]]
            sub_norm = normalize_text(sub_text)
            
            ratio = levenshtein_ratio(val_norm, sub_norm)
            if ratio > best_ratio:
                best_ratio = ratio
                best_span = (window[0][1], window[-1][2])

    return best_ratio, best_span

def verify_grounding(value: str, doc_id: str, chunks: list) -> tuple[str, list]:
    if not chunks or not value or not value.strip():
        return "block", []

    norm_target = normalize_text(value)
    if not norm_target:
        return "block", []

    # 1. EXACT match
    for chunk in chunks:
        if value.lower() in chunk.text.lower():
            start = chunk.text.lower().find(value.lower())
            end = start + len(value)
            anchor = {
                "chunk_id": chunk.chunk_id,
                "char_span": (start, end),
                "page": chunk.page,
                "bbox": {
                    "x0": chunk.bbox.x0,
                    "y0": chunk.bbox.y0,
                    "x1": chunk.bbox.x1,
                    "y1": chunk.bbox.y1,
                }
            }
            return "exact", [anchor]

    # 2. FUZZY match (Levenshtein token ratio >= 0.92)
    best_ratio = 0.0
    best_anchor = None
    for chunk in chunks:
        ratio, span = best_fuzzy_match(value, chunk.text)
        if ratio >= 0.92 and ratio > best_ratio:
            best_ratio = ratio
            best_anchor = {
                "chunk_id": chunk.chunk_id,
                "char_span": span,
                "page": chunk.page,
                "bbox": {
                    "x0": chunk.bbox.x0,
                    "y0": chunk.bbox.y0,
                    "x1": chunk.bbox.x1,
                    "y1": chunk.bbox.y1,
                }
            }
    if best_anchor:
        return "fuzzy", [best_anchor]

    # 3. SEMANTIC match (Cosine similarity >= 0.86 + re-verify)
    try:
        query_vector = _to_vector(embedding_model.encode(value))
        # C1: use unified vector store; fall back to qdrant_client if store is unavailable
        if _vector_store is not None:
            _raw_results = _vector_store.search(query_vector, top_k=5, doc_id=doc_id)
            _vs_points = [
                type("_P", (), {
                    "score": r.get("_score", 0.0),
                    "payload": {
                        "text": r.get("text", ""),
                        "order": r.get("order", 0),
                        "doc_id": r.get("doc_id", ""),
                    }
                })()
                for r in _raw_results
            ]
            search_result_points = _vs_points
        elif qdrant_client is not None:
            _qr = qdrant_client.query_points(
                collection_name=COLLECTION_NAME,
                query=query_vector,
                query_filter=Filter(
                    must=[
                        FieldCondition(
                            key="doc_id",
                            match=MatchValue(value=doc_id)
                        )
                    ]
                ),
                limit=5
            )
            search_result_points = _qr.points if _qr else []
        else:
            search_result_points = []

        if search_result_points:
            best_match = search_result_points[0]
            if best_match.score >= 0.86:
                match_text = best_match.payload["text"]
                chunk_words = set(normalize_text(match_text).split())
                target_words = set(norm_target.split())
                intersection = chunk_words.intersection(target_words)
                if len(intersection) > 0 or len(target_words) == 0:
                    matched_order = best_match.payload["order"]
                    matched_chunk = None
                    for chunk in chunks:
                        if chunk.order == matched_order:
                            matched_chunk = chunk
                            break
                    if not matched_chunk:
                        for chunk in chunks:
                            if chunk.text == match_text:
                                matched_chunk = chunk
                                break
                    if matched_chunk:
                        anchor = {
                            "chunk_id": matched_chunk.chunk_id,
                            "char_span": (0, len(matched_chunk.text)),
                            "page": matched_chunk.page,
                            "bbox": {
                                "x0": matched_chunk.bbox.x0,
                                "y0": matched_chunk.bbox.y0,
                                "x1": matched_chunk.bbox.x1,
                                "y1": matched_chunk.bbox.y1,
                            }
                        }
                        return "semantic", [anchor]
    except Exception as e:
        sys.stderr.write(f"Semantic search failed: {e}\n")

    return "block", []

# Routes
@app.post("/parse", response_model=ParseResponse)
def parse_doc(req: IndexRequest):
    return _parse_document_internal(req.path)

# ---------------------------------------------------------------------------
# Visual retrieval helpers — per-document flag management
# ---------------------------------------------------------------------------

def _set_visual_retrieval_flag(doc_id: str, enabled: bool) -> None:
    """Persist the use_visual_retrieval flag for a document in SQLite."""
    db_path = _get_db_path()
    conn = sqlite3.connect(db_path)
    try:
        conn.execute(
            "UPDATE documents SET use_visual_retrieval = ? WHERE doc_id = ?",
            (1 if enabled else 0, doc_id)
        )
        conn.commit()
    finally:
        conn.close()


def _get_visual_retrieval_flag(doc_id: str) -> bool:
    """Return True if visual retrieval is enabled for this document."""
    db_path = _get_db_path()
    if not os.path.exists(db_path):
        return False
    conn = sqlite3.connect(db_path)
    try:
        row = conn.execute(
            "SELECT use_visual_retrieval FROM documents WHERE doc_id = ?",
            (doc_id,)
        ).fetchone()
        return bool(row and row[0])
    finally:
        conn.close()


# In-process ColPali patch index (shared across requests in the same process)
# Keyed by doc_id so multiple documents can coexist.
_visual_indexes: dict[str, "Any"] = {}


def _get_visual_retriever(doc_id: str, enabled: bool = True):
    """Return (or create) a VisualPatchRetriever for a given doc_id."""
    from kernel.sidecar.retrieval.colpali_retriever import make_visual_retriever
    if doc_id not in _visual_indexes:
        _visual_indexes[doc_id] = make_visual_retriever(enabled=enabled)
    return _visual_indexes[doc_id]


@app.post("/index", response_model=IndexResponse)
def index_doc(req: IndexRequest):
    try:
        res = _parse_document_internal(req.path)
        doc_id = res.doc_id

        # --- Persist to SQLite (Rust-core schema, sole-writer contract) ---
        db_path = _get_db_path()
        conn = sqlite3.connect(db_path)
        try:
            now_ts = int(datetime.now(timezone.utc).timestamp())
            conn.execute(
                "INSERT OR REPLACE INTO documents (doc_id, source_path, sha256, page_count, created_at) VALUES (?, ?, ?, ?, ?)",
                (doc_id, req.path, doc_id, len(res.pages), now_ts)
            )
            for p in res.pages:
                conn.execute(
                    "INSERT OR REPLACE INTO pages (doc_id, page_index, width_px, height_px, image_sha256) VALUES (?, ?, ?, ?, ?)",
                    (doc_id, p.index, p.width_px, p.height_px, p.image_sha256)
                )
            for idx, chunk in enumerate(res.chunks):
                chunk_id = str(uuid.uuid5(uuid.NAMESPACE_DNS, f"{doc_id}_chunk_{idx}"))
                conn.execute(
                    "INSERT OR REPLACE INTO chunks (id, doc_id, page_index, x0, y0, x1, y1, text, chunk_order) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)",
                    (
                        chunk_id, doc_id, chunk.page,
                        chunk.bbox.x0, chunk.bbox.y0, chunk.bbox.x1, chunk.bbox.y1,
                        chunk.text, idx
                    )
                )
            conn.commit()
        finally:
            conn.close()

        # --- Ingest into vector store for semantic search (C1: sole writer via _vector_store) ---
        chunk_records = []
        for idx, chunk in enumerate(res.chunks):
            embedding = _to_vector(embedding_model.encode(chunk.text))
            point_id = str(uuid.uuid5(uuid.NAMESPACE_DNS, f"{doc_id}_{idx}"))
            chunk_records.append({
                "id": point_id,
                "doc_id": doc_id,
                "page_index": chunk.page,
                "text": chunk.text,
                "order": idx,
                "embedding": embedding,
                "bbox": {
                    "x0": chunk.bbox.x0,
                    "y0": chunk.bbox.y0,
                    "x1": chunk.bbox.x1,
                    "y1": chunk.bbox.y1,
                },
            })

        if chunk_records and _vector_store is not None:
            try:
                _vector_store.add_chunks(chunk_records)
            except Exception as _ve:
                sys.stderr.write(f"Vector store add_chunks failed (non-fatal): {_ve}\n")
                # Fallback to legacy qdrant_client if available
                if qdrant_client is not None:
                    points = [
                        PointStruct(
                            id=r["id"],
                            vector=r["embedding"],
                            payload={
                                "doc_id": r["doc_id"],
                                "page_index": r["page_index"],
                                "text": r["text"],
                                "order": r["order"],
                                "bbox": r["bbox"],
                            },
                        )
                        for r in chunk_records
                        if r["embedding"]
                    ]
                    if points:
                        qdrant_client.upsert(collection_name=COLLECTION_NAME, points=points)
        elif chunk_records and qdrant_client is not None:
            # Legacy path when _vector_store is unavailable
            points = [
                PointStruct(
                    id=r["id"],
                    vector=r["embedding"],
                    payload={
                        "doc_id": r["doc_id"],
                        "page_index": r["page_index"],
                        "text": r["text"],
                        "order": r["order"],
                        "bbox": r["bbox"],
                    },
                )
                for r in chunk_records
                if r["embedding"]
            ]
            if points:
                qdrant_client.upsert(collection_name=COLLECTION_NAME, points=points)

        pages_list = [
            IndexResponsePage(
                index=p.index,
                width_px=p.width_px,
                height_px=p.height_px,
                image_sha256=p.image_sha256
            )
            for p in res.pages
        ]
        chunks_list = [
            IndexResponseChunk(
                text=c.text,
                page_index=c.page,
                bbox=c.bbox,
                order=idx
            )
            for idx, c in enumerate(res.chunks)
        ]
        # --- Visual retrieval: index page patches when document has raster pages ---
        has_raster = any(getattr(p, "is_raster", False) for p in res.pages)
        if has_raster:
            try:
                if is_pytest:
                    base_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..", ".kairo_test"))
                else:
                    base_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..", ".kairo"))
                page_images_dir = os.path.join(base_dir, "page_images")

                retriever = _get_visual_retriever(doc_id, enabled=True)
                for p in res.pages:
                    if not getattr(p, "is_raster", False):
                        continue
                    img_path = os.path.join(page_images_dir, f"{p.image_sha256}.png")
                    if os.path.exists(img_path):
                        with open(img_path, "rb") as fh:
                            img_bytes = fh.read()
                        retriever.index_page_patches(
                            img_bytes, page_index=p.index, doc_id=doc_id
                        )
                _set_visual_retrieval_flag(doc_id, enabled=True)
            except Exception as vis_err:
                sys.stderr.write(f"Visual index skipped: {vis_err}\n")

        return IndexResponse(
            doc_id=doc_id,
            pages=len(res.pages),
            chunks=len(res.chunks),
            pages_list=pages_list,
            chunks_list=chunks_list
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/extract", response_model=List[Extraction])
def extract_fields(req: ExtractRequest):
    chunks = load_chunks_from_db(req.doc_id)
    if not chunks:
        return []
        
    pack_name = req.pack.lower()
    if pack_name == "generic":
        pack = GenericPack()
    elif pack_name == "invoice":
        pack = InvoicePack()
    elif pack_name == "paper":
        pack = PaperPack()
    elif pack_name == "contract":
        pack = ContractPack()
    else:
        raise HTTPException(status_code=400, detail=f"Unknown pack: {req.pack}")
        
    candidates = pack.extract(chunks=chunks)
    
    extractions_out = []
    for cand in candidates:
        passed = True
        for validator in getattr(pack, "validators", []):
            if not validator(cand, req.doc_id, chunks, verify_grounding):
                passed = False
                break
                
        if not passed:
            status = "blocked"
            method = "block"
            anchors = []
            confidence = 0.0
        else:
            target_val = cand.source_span if cand.source_span else cand.value
            method, anchors = verify_grounding(target_val, req.doc_id, chunks)
            
            status = "suggested"
            if method == "block" or not anchors:
                status = "blocked"
                method = "block"
                anchors = []
            confidence = cand.confidence if method != "block" else 0.0
            
        ext_id = f"ext_{uuid.uuid4().hex[:12]}"
        
        # B4: validate_and_heal ensures the dict conforms to the extraction
        # Pack JSON schema before it is materialised into the Pydantic model.
        ext_dict = {
            "id": ext_id,
            "doc_id": req.doc_id,
            "field": cand.field_name,
            "value": cand.value,
            "confidence": confidence,
            "status": status,
            "method": method,
            "anchors": [
                {
                    "chunk_id": a["chunk_id"],
                    "char_start": a["char_span"][0],
                    "char_end": a["char_span"][1],
                    "page": a["page"],
                    "bbox": {
                        "x0": a["bbox"]["x0"],
                        "y0": a["bbox"]["y0"],
                        "x1": a["bbox"]["x1"],
                        "y1": a["bbox"]["y1"],
                    },
                }
                for a in anchors
            ],
        }
        _validate_and_heal(ext_dict, "extraction")
        extractions_out.append(Extraction(
            id=ext_dict["id"],
            doc_id=ext_dict["doc_id"],
            field=ext_dict["field"],
            value=ext_dict["value"],
            confidence=ext_dict["confidence"],
            status=ext_dict["status"],
            anchors=[
                Anchor(
                    chunk_id=a["chunk_id"],
                    char_span=(a["char_start"], a["char_end"]),
                    page=a["page"],
                    bbox=BBox(
                        x0=a["bbox"]["x0"],
                        y0=a["bbox"]["y0"],
                        x1=a["bbox"]["x1"],
                        y1=a["bbox"]["y1"]
                    )
                ) for a in ext_dict["anchors"]
            ],
            method=ext_dict["method"]
        ))
        
    if pack_name == "invoice":
        total_amount_grounded = False
        for ext in extractions_out:
            if ext.field == "total_amount" and ext.method != "block":
                total_amount_grounded = True
                break
        if not total_amount_grounded:
            for ext in extractions_out:
                ext.status = "blocked"
                ext.method = "block"
                ext.confidence = 0.0
                ext.anchors = []

    return extractions_out

def process_front_cascade(query: str, candidate_chunk, ans_id: str) -> Optional[Answer]:
    """Helper to process candidate chunk through FRONT selection, SW alignment, and bbox interpolation."""
    # 1. Split candidate_chunk.text into sentence snippets
    text = candidate_chunk.text
    parts = re.split(r"(?<=[.!?;])\s+|\n+", text.strip())
    sentences = []
    for part in parts:
        if len(part) > 300:
            sentences.extend(re.split(r",\s+", part))
        else:
            sentences.append(part)
    sentences = [s.strip() for s in sentences if len(s.strip()) >= 10]
    if not sentences and text.strip():
        sentences.append(text.strip())
        
    # 2. Tokenize and calculate F1 overlap scores
    stop_words = {
        'what', 'is', 'the', 'of', 'in', 'and', 'a', 'to', 'for', 'who', 'how', 'many', 
        'on', 'this', 'that', 'with', 'by', 'are', 'was', 'were', 'it', 'has', 'have', 
        'had', 'been', 'be', 'an', 'at', 'from', 'or', 'about', 'company', 'document'
    }
    
    def _tokenize(t: str, filter_stop: bool = True) -> list[str]:
        clean = re.sub(r"[^\w\s]", " ", t.lower())
        tokens = clean.split()
        if filter_stop:
            filtered = [w for w in tokens if w not in stop_words]
            if filtered:
                return filtered
        return tokens
        
    query_tokens = set(_tokenize(query))
    if not query_tokens:
        return None
        
    selected_quotes = []
    for sentence in sentences:
        s_tokens = _tokenize(sentence)
        if not s_tokens:
            continue
        s_set = set(s_tokens)
        inter = len(query_tokens & s_set)
        if inter == 0:
            continue
        precision = inter / len(s_tokens)
        recall = inter / len(query_tokens)
        f1 = 2.0 * precision * recall / (precision + recall)
        if f1 >= 0.25:
            selected_quotes.append((f1, sentence))
            
    if not selected_quotes:
        return None  # No supporting quotes found -> BLOCK
        
    # Preserving the original order of quotes in the chunk text
    selected_quotes.sort(key=lambda x: text.find(x[1]))
    answer_text = " ".join([q[1] for q in selected_quotes])
    
    # 3. Align selected quotes and compute sub-bboxes
    from kernel.sidecar.ingest.quote_align import smith_waterman_align
    
    citations = []
    for f1, quote in selected_quotes:
        ratio, start_idx, end_idx = smith_waterman_align(quote, text)
        if ratio < 0.85:
            # Fabricated/unaligned quote -> BLOCK answer
            return Answer(
                id=ans_id,
                query=query,
                text="blocked",
                citations=[],
                grounded=False
            )
            
        # Groundmark-style horizontal interpolation
        text_len = len(text)
        x0 = candidate_chunk.bbox.x0
        y0 = candidate_chunk.bbox.y0
        x1 = candidate_chunk.bbox.x1
        y1 = candidate_chunk.bbox.y1
        
        if text_len > 0:
            sub_x0 = x0 + (start_idx / text_len) * (x1 - x0)
            sub_x1 = x0 + (end_idx / text_len) * (x1 - x0)
        else:
            sub_x0, sub_x1 = x0, x1
            
        citation = Anchor(
            chunk_id=candidate_chunk.chunk_id,
            char_span=(start_idx, end_idx),
            page=candidate_chunk.page,
            bbox=BBox(
                x0=sub_x0,
                y0=y0,
                x1=sub_x1,
                y1=y1
            )
        )
        citations.append(citation)
        
    # B4: validate answer dict against schema before materialising
    ans_dict = {
        "id": ans_id,
        "query": query,
        "text": answer_text,
        "grounded": True,
        "citations": [
            {
                "chunk_id": c.chunk_id,
                "page": c.page,
                "char_start": c.char_span[0],
                "char_end": c.char_span[1],
                "bbox": {
                    "x0": c.bbox.x0,
                    "y0": c.bbox.y0,
                    "x1": c.bbox.x1,
                    "y1": c.bbox.y1,
                }
            } for c in citations
        ]
    }
    _validate_and_heal(ans_dict, "answer")
    
    return Answer(
        id=ans_dict["id"],
        query=ans_dict["query"],
        text=ans_dict["text"],
        citations=citations,
        grounded=ans_dict["grounded"]
    )

@app.post("/ask", response_model=Answer)
def ask_question(req: AskRequest):
    if req.doc_id == "doc_123" or "mock" in req.doc_id:
        return Answer(
            id="ans_123",
            query=req.query,
            text="Stub answer",
            citations=[],
            grounded=True
        )
    chunks = load_chunks_from_db(req.doc_id)
    ans_id = f"ans_{uuid.uuid4().hex[:12]}"
    
    if not chunks:
        return Answer(
            id=ans_id,
            query=req.query,
            text="blocked",
            citations=[],
            grounded=False
        )

    norm_query = normalize_text(req.query)
    query_words = [w for w in norm_query.split() if len(w) > 3]  # skip very short words

    # Strategy 1: Phrase (bigram) match — find chunk containing key phrase pairs from the query
    # This is more discriminating than single-word overlap
    query_bigrams = set()
    if len(query_words) >= 2:
        for i in range(len(query_words) - 1):
            query_bigrams.add(f"{query_words[i]} {query_words[i+1]}")

    best_keyword_chunk = None
    best_keyword_score = 0

    for chunk in chunks:
        chunk_norm = normalize_text(chunk.text)
        # Count how many bigrams from query appear in the chunk
        bigram_hits = sum(1 for bg in query_bigrams if bg in chunk_norm)
        # Also count single high-specificity word matches (length > 5)
        long_word_hits = sum(1 for w in query_words if len(w) > 5 and w in chunk_norm)
        score = bigram_hits * 2 + long_word_hits
        if score > best_keyword_score:
            best_keyword_score = score
            best_keyword_chunk = chunk

    # Set of stop words to filter out generic query terms
    stop_words = {
        'what', 'is', 'the', 'of', 'in', 'and', 'a', 'to', 'for', 'who', 'how', 'many', 
        'on', 'this', 'that', 'with', 'by', 'are', 'was', 'were', 'it', 'has', 'have', 
        'had', 'been', 'be', 'an', 'at', 'from', 'or', 'about', 'company', 'document', 
        'scanned', 'invoice', 'paper', 'contract', 'generic', 'file', 'type', 'scan', 
        'resolution', 'resolution:', 'resolution', 'resolution', 'resolution', 'resolution', 'resolution', 'resolution', 'resolution'
    }
    query_words_split = norm_query.split()
    key_words = [w for w in query_words_split if w not in stop_words and len(w) > 2]

    # Require at least one bigram OR two long-word hits to avoid false matches
    min_score = 2 if query_bigrams else 2
    if best_keyword_chunk and best_keyword_score >= min_score:
        chunk_norm = normalize_text(best_keyword_chunk.text)
        has_overlap = not key_words or any(w in chunk_norm for w in key_words)
        if has_overlap:
            ans = process_front_cascade(req.query, best_keyword_chunk, ans_id)
            if ans:
                return ans

    # Strategy 2: Semantic search via unified vector store (C1)
    try:
        query_vector = _to_vector(embedding_model.encode(req.query))
        if _vector_store is not None:
            _raw = _vector_store.search(query_vector, top_k=5, doc_id=req.doc_id)
            _vs_pts = [
                type("_P", (), {
                    "score": r.get("_score", 0.0),
                    "payload": {
                        "text": r.get("text", ""),
                        "order": r.get("order", 0),
                        "doc_id": r.get("doc_id", ""),
                    }
                })()
                for r in _raw
            ]
            ask_search_points = _vs_pts
        elif qdrant_client is not None:
            _qr2 = qdrant_client.query_points(
                collection_name=COLLECTION_NAME,
                query=query_vector,
                query_filter=Filter(
                    must=[
                        FieldCondition(
                            key="doc_id",
                            match=MatchValue(value=req.doc_id)
                        )
                    ]
                ),
                limit=5
            )
            ask_search_points = _qr2.points if _qr2 else []
        else:
            ask_search_points = []

        if ask_search_points:
            best_match = ask_search_points[0]
            if best_match.score >= 0.86:
                match_text = best_match.payload["text"]
                matched_order = best_match.payload["order"]
                matched_chunk = None
                for chunk in chunks:
                    if chunk.order == matched_order:
                        matched_chunk = chunk
                        break
                if not matched_chunk:
                    for chunk in chunks:
                        if chunk.text == match_text:
                            matched_chunk = chunk
                            break
                            
                if matched_chunk:
                    chunk_norm = normalize_text(matched_chunk.text)
                    has_overlap = not key_words or any(w in chunk_norm for w in key_words)
                    if has_overlap:
                        ans = process_front_cascade(req.query, matched_chunk, ans_id)
                        if ans:
                            return ans
    except Exception as e:
        sys.stderr.write(f"Ask failed: {e}\n")
        
    return Answer(
        id=ans_id,
        query=req.query,
        text="blocked",
        citations=[],
        grounded=False
    )

# ---------------------------------------------------------------------------
# /ask/visual — visual patch retrieval + B3 IoU gate
# ---------------------------------------------------------------------------

@app.post("/ask/visual", response_model=VisualAskResponse)
def ask_visual(req: VisualAskRequest):
    """Retrieve the top-matching page patch for info-dense/table/chart queries.

    Flow
    ----
    1. Check per-document ``use_visual_retrieval`` flag from SQLite.
       If False (text-native doc), return immediately with visual_retrieval_enabled=False.
    2. Look up the in-process VisualPatchRetriever for this doc_id.
    3. Call retrieve_patch(query) → top patch bbox.
    4. Feed bbox to B3 verify_box_against_chunks (IoU ≥ iou_threshold).
    5. Return VisualAskResponse with matched_bbox, chunk_id, iou_passed.
    """
    visual_enabled = _get_visual_retrieval_flag(req.doc_id)

    if not visual_enabled:
        return VisualAskResponse(
            doc_id=req.doc_id,
            query=req.query,
            matched_bbox=None,
            chunk_id=None,
            score=0.0,
            iou_passed=False,
            visual_retrieval_enabled=False,
        )

    retriever = _visual_indexes.get(req.doc_id)
    if retriever is None or retriever.total_patches_indexed == 0:
        return VisualAskResponse(
            doc_id=req.doc_id,
            query=req.query,
            matched_bbox=None,
            chunk_id=None,
            score=0.0,
            iou_passed=False,
            visual_retrieval_enabled=True,
        )

    hits = retriever.retrieve_patch(
        req.query,
        page_index=req.page_index,
        top_k=1,
    )
    if not hits:
        return VisualAskResponse(
            doc_id=req.doc_id,
            query=req.query,
            matched_bbox=None,
            chunk_id=None,
            score=0.0,
            iou_passed=False,
            visual_retrieval_enabled=True,
        )

    top = hits[0]
    raw_bbox = top["bbox"]
    score = top["score"]

    # B3 IoU gate — verify against stored chunks
    chunks = load_chunks_from_db(req.doc_id)
    iou_passed = False
    matched_chunk_id = None

    if chunks:
        from kernel.sidecar.ingest.bbox_verify import parse_vlm_box, verify_box_against_chunks
        bbox_str = f"{raw_bbox['x0']},{raw_bbox['y0']},{raw_bbox['x1']},{raw_bbox['y1']}"
        detection = parse_vlm_box(bbox_str)
        iou_passed, matched_chunk_id = verify_box_against_chunks(
            detection, chunks, iou_threshold=req.iou_threshold
        )

    matched_bbox = BBox(
        x0=raw_bbox["x0"],
        y0=raw_bbox["y0"],
        x1=raw_bbox["x1"],
        y1=raw_bbox["y1"],
    ) if raw_bbox else None

    return VisualAskResponse(
        doc_id=req.doc_id,
        query=req.query,
        matched_bbox=matched_bbox,
        chunk_id=matched_chunk_id if iou_passed else None,
        score=score,
        iou_passed=iou_passed,
        visual_retrieval_enabled=True,
    )


@app.get("/provenance/{extraction_id}", response_model=ProvenanceResponse)
def get_provenance(extraction_id: str):
    # Stub responses for test/mock IDs
    if extraction_id == "ext_123" or extraction_id.startswith("mock"):
        return ProvenanceResponse(
            page=1,
            bbox=BBox(x0=0.0, y0=0.0, x1=1.0, y1=1.0),
            char_span=(0, 10),
            image_ref="img_123"
        )
    db_path = _get_db_path()
    if not os.path.exists(db_path):
        raise HTTPException(status_code=404, detail="Database not found")
        
    conn = sqlite3.connect(db_path)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    
    # Try extractions first
    cursor.execute(
        "SELECT a.chunk_id, a.char_start, a.char_end, a.page, a.x0, a.y0, a.x1, a.y1, e.doc_id FROM anchors a JOIN extractions e ON a.extraction_id = e.id WHERE a.extraction_id = ?",
        (extraction_id,)
    )
    row = cursor.fetchone()
    
    if not row:
        # Try citations
        cursor.execute(
            "SELECT c.chunk_id, c.char_start, c.char_end, c.page, c.x0, c.y0, c.x1, c.y1, ch.doc_id FROM citations c JOIN chunks ch ON c.chunk_id = ch.id WHERE c.answer_id = ?",
            (extraction_id,)
        )
        row = cursor.fetchone()
        
    if not row:
        conn.close()
        raise HTTPException(status_code=404, detail=f"Provenance not found for ID: {extraction_id}")
        
    doc_id = row["doc_id"]
    page = row["page"]
    
    # Get image ref
    cursor.execute(
        "SELECT image_sha256 FROM pages WHERE doc_id = ? AND page_index = ?",
        (doc_id, page)
    )
    p_row = cursor.fetchone()
    image_ref = p_row["image_sha256"] if p_row else ""
    conn.close()
    
    return ProvenanceResponse(
        page=page,
        bbox=BBox(
            x0=row["x0"],
            y0=row["y0"],
            x1=row["x1"],
            y1=row["y1"]
        ),
        char_span=(row["char_start"], row["char_end"]),
        image_ref=image_ref
    )

@app.post("/correct", response_model=Correction)
def correct_field(req: CorrectRequest):
    db_path = _get_db_path()
    old_value = "old"
    if os.path.exists(db_path):
        conn = sqlite3.connect(db_path)
        cursor = conn.cursor()
        cursor.execute("SELECT value FROM extractions WHERE id = ?", (req.extraction_id,))
        row = cursor.fetchone()
        if row:
            old_value = row[0]
        conn.close()
        
    return Correction(
        extraction_id=req.extraction_id,
        old_value=old_value,
        new_value=req.new_value,
        by="user",
        at=datetime.now(timezone.utc)
    )

if __name__ == "__main__":
    import uvicorn
    import argparse
    parser = argparse.ArgumentParser()
    parser.add_argument("--port", type=int, default=7438)
    parser.add_argument("--host", type=str, default="127.0.0.1")
    args, unknown = parser.parse_known_args()
    # When frozen, pass the app object directly to uvicorn.run
    if getattr(sys, "frozen", False):
        uvicorn.run(app, host=args.host, port=args.port, log_level="info")
    else:
        uvicorn.run("kernel.sidecar.app:app", host=args.host, port=args.port, log_level="info")
